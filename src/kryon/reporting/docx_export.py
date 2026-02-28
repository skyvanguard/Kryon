"""DOCX export — convert HTML reports to Microsoft Word format."""

from __future__ import annotations

import logging
import re
from io import BytesIO

logger = logging.getLogger(__name__)


def html_to_docx(html: str) -> bytes:
    """Convert HTML report to DOCX bytes using python-docx.

    Requires: pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

    doc = Document()

    # Strip HTML tags for plain text extraction
    text = _strip_html(html)

    # Parse sections from HTML
    sections = _extract_sections(html)

    for section in sections:
        if section["type"] == "h1":
            p = doc.add_heading(section["text"], level=1)
        elif section["type"] == "h2":
            p = doc.add_heading(section["text"], level=2)
        elif section["type"] == "h3":
            p = doc.add_heading(section["text"], level=3)
        elif section["type"] == "table":
            _add_table(doc, section["rows"])
        elif section["type"] == "p":
            if section["text"].strip():
                doc.add_paragraph(section["text"])

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _strip_html(html: str) -> str:
    """Strip HTML tags and decode entities."""
    text = re.sub(r"<[^>]+>", "", html)
    text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    text = text.replace("&rarr;", "->").replace("&mdash;", "-")
    return text.strip()


def _extract_sections(html: str) -> list[dict]:
    """Extract document sections from HTML."""
    sections = []

    # Find headings
    for match in re.finditer(r"<(h[1-3])[^>]*>(.*?)</\1>", html, re.DOTALL | re.IGNORECASE):
        tag, text = match.group(1).lower(), _strip_html(match.group(2))
        sections.append({"type": tag, "text": text, "pos": match.start()})

    # Find paragraphs
    for match in re.finditer(r"<p[^>]*>(.*?)</p>", html, re.DOTALL | re.IGNORECASE):
        text = _strip_html(match.group(1))
        sections.append({"type": "p", "text": text, "pos": match.start()})

    # Find tables
    for match in re.finditer(r"<table[^>]*>(.*?)</table>", html, re.DOTALL | re.IGNORECASE):
        rows = _parse_table(match.group(1))
        sections.append({"type": "table", "rows": rows, "pos": match.start()})

    # Sort by position
    sections.sort(key=lambda s: s.get("pos", 0))

    # If no sections found, treat whole text as one paragraph
    if not sections:
        text = _strip_html(html)
        if text:
            sections.append({"type": "p", "text": text})

    return sections


def _parse_table(table_html: str) -> list[list[str]]:
    """Parse HTML table into list of rows."""
    rows = []
    for tr in re.finditer(r"<tr[^>]*>(.*?)</tr>", table_html, re.DOTALL | re.IGNORECASE):
        cells = re.findall(r"<t[hd][^>]*>(.*?)</t[hd]>", tr.group(1), re.DOTALL | re.IGNORECASE)
        rows.append([_strip_html(c) for c in cells])
    return rows


def _add_table(doc, rows: list[list[str]]) -> None:
    """Add a table to the DOCX document."""
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < max_cols:
                table.rows[i].cells[j].text = cell_text
